from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('prof_pic.jpg', width = Inches(2.0))

# name, phone number and email details
speak('What is your name?')
name = input('Name : ')
speak('Hello ' + name + ' how are you today?')

speak('What is your phone number?')
phone_number = input('Phone Number : ')
speak('Your phone number is ' + str(phone_number) + '. Is it correct?')
correct_phone_number = input("Yes or No : ")

while True:
    if correct_phone_number.upper() == 'NO':
        phone_number = input('Phone Number : ')
        speak('Your phone number is ' + str(phone_number) + '. Is it correct?')
        correct_phone_number = input("Yes or No ")
    else:
        break

speak('Enter your email address')
email = input('eMail : ')

document.add_paragraph(name + ' | ' + phone_number + '|' + email)

# about me
speak('Tell us about yourself')
document.add_heading('About me')
about_me = input('Yourself : ')
document.add_paragraph(about_me)

# work experience
speak('Work experience')
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Enter company')
company = input('Company : ')
speak('Date from')
from_date = input('From Date : ')
speak('Date to')
to_date = input('To Date : ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Your experience at ' + company)
experience_details = input('Describe your experience at ' + company + ' : ')
p.add_run(experience_details)

# more experiences
while True:
    speak('Do you have more experiences?')
    has_more_experiences = input('Yes or No : ')
    if has_more_experiences.upper() == 'YES':
        p = document.add_paragraph()

        speak('Enter company')
        company = input('Company : ')
        speak('Date from')
        from_date = input('From Date : ')
        speak('Date to')
        to_date = input('To Date : ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        speak('Your experience at ' + company)
        experience_details = input('Describe your experience at ' + company + ' : ')
        p.add_run(experience_details)
    else:
        break

# skiils
document.add_heading('Skills')
speak('Enter skills')
skill = input('Enter skill : ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# more skills
while True:
    speak('Do you have more skills?')
    has_more_skills = input('Yes or No : ')
    if has_more_skills.upper() == 'YES':
        speak('Enter skill')
        skill = input('Enter skill : ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
p = document.sections[0].footer.paragraphs[0]
p.text = 'CV generated using Amigoscode and Intuit Quickbooks'

document.save('cv.docx')